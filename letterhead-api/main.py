from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image
import base64
import io
import os
import subprocess

app = FastAPI()

# ✅ Allow frontend origin from Vercel
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://letterheadcreation.vercel.app"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ✅ Ensure temp folder exists
TEMP_FOLDER = "temp_images"
os.makedirs(TEMP_FOLDER, exist_ok=True)

# ✅ Split header image
def split_letterhead_image(image_bytes):
    image = Image.open(io.BytesIO(image_bytes))
    width, height = image.size
    header_img = image.crop((0, 0, width, int(height * 0.25)))
    header_io = io.BytesIO()
    header_img.save(header_io, format="PNG")
    header_io.seek(0)
    return header_io

# ✅ Convert HTML to image using wkhtmltoimage
def render_footer_html_to_image(html_string, output_path):
    html_file = os.path.join(TEMP_FOLDER, "footer_temp.html")
    with open(html_file, "w", encoding="utf-8") as f:
        f.write(f"<html><body><div style='font-family:Arial; font-size:12pt;'>{html_string}</div></body></html>")

    try:
        subprocess.run(
            ["wkhtmltoimage", "--width", "800", "--height", "100", html_file, output_path],
            check=True
        )
    except subprocess.CalledProcessError as e:
        print("❌ wkhtmltoimage error:", e)

@app.post("/merge-docx/")
async def merge_docx(request: Request):
    try:
        data = await request.json()
        base64_data = data["image"].split(";base64,")[-1]
        content = data.get("content", "")
        footer_html = data.get("footer_text", "")

        image_bytes = base64.b64decode(base64_data)
        header_stream = split_letterhead_image(image_bytes)

        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # ✅ Add header
        header_para = section.header.paragraphs[0]
        header_para.alignment = 1
        header_para.add_run().add_picture(header_stream, width=Inches(6.5))

        # ✅ Render and add footer image
        if footer_html.strip():
            footer_img_path = os.path.join(TEMP_FOLDER, "footer_rendered.png")
            render_footer_html_to_image(footer_html, footer_img_path)

            if os.path.exists(footer_img_path):
                with open(footer_img_path, "rb") as f:
                    section.footer.paragraphs[0].add_run().add_picture(f, width=Inches(6.5))

        # ✅ Add body
        for line in content.split("\n"):
            if line.strip():
                para = doc.add_paragraph(line.strip())
                para.style.font.size = Pt(11)

        # ✅ Return document
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=letterhead_final.docx"}
        )

    except Exception as e:
        print("❌ Error generating document:", str(e))
        return JSONResponse(status_code=500, content={"error": str(e)})
